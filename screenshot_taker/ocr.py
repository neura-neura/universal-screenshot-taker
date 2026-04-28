from __future__ import annotations

import base64
import html
import importlib.metadata
import importlib.util
import json
import locale
import os
import re
import shutil
import subprocess
import sys
import tempfile
import time
import textwrap
import unicodedata
import urllib.error
import urllib.request
from dataclasses import dataclass, field
from difflib import SequenceMatcher
from pathlib import Path
from typing import Callable, Iterable

from PIL import Image
from pypdf import PdfReader, PdfWriter
from reportlab.lib.enums import TA_JUSTIFY
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.utils import ImageReader
from reportlab.pdfgen import canvas
from reportlab.platypus import Image as PlatypusImage, PageBreak, Paragraph, SimpleDocTemplate, Spacer

from .capture import ensure_directory, export_pdf


ProgressCallback = Callable[[int, int, str], None]
StatusCallback = Callable[[str], None]

DEFAULT_OLLAMA_URL = "http://localhost:11434"
DEFAULT_OCR_PROMPT = (
    "You are performing OCR on a book page screenshot. Extract only the text that is visible in the image. "
    "Preserve reading order, visible line breaks, and paragraph spacing as closely as possible. "
    "Preserve real paragraph breaks. If a new paragraph clearly starts on this page, keep a blank line before it. "
    "If the page begins or ends in the middle of an existing paragraph, continue that paragraph naturally instead of inventing a new one. "
    "If this page contains a chart, diagram, table, illustration, equation block, or other visual content that should be preserved as an image in rich exports, "
    "add a separate line exactly [[KEEP_PAGE_IMAGE]] near the top of your response. "
    "Ignore running headers, repeated headers, footers, page numbers, marginal page ornaments, and other decorative or repeated page furniture "
    "unless they are clearly part of the main body text. "
    "Keep one output line per visual text line whenever possible. Return plain text only. "
    "Do not add explanations, Markdown fences, comments, HTML, or XML tags."
)
OLLAMA_WINDOWS_INSTALL_COMMAND = "irm https://ollama.com/install.ps1 | iex"
OLLAMA_WINDOWS_RUN_COMMAND = "ollama serve"
TESSERACT_WINDOWS_INSTALL_COMMAND = "winget install -e --id UB-Mannheim.TesseractOCR"
OCRMYPDF_WINDOWS_INSTALL_COMMAND = ".\\venv\\Scripts\\python.exe -m pip install ocrmypdf"
SURYA_INSTALL_COMMAND = (
    ".\\venv\\Scripts\\python.exe -m pip install --upgrade "
    "\"surya-ocr\" \"transformers>=4.56.1,<5\""
)
TESSERACT_TESSDATA_FAST_BASE_URL = "https://raw.githubusercontent.com/tesseract-ocr/tessdata_fast/main"

TESSERACT_LANGUAGE_PRESETS = (
    ("eng", "tess_lang_english"),
    ("spa", "tess_lang_spanish"),
    ("chi_sim", "tess_lang_chinese_simplified"),
    ("chi_tra", "tess_lang_chinese_traditional"),
    ("eng+spa", "tess_lang_english_spanish"),
    ("eng+spa+chi_sim", "tess_lang_english_spanish_chinese_simplified"),
    ("fra", "tess_lang_french"),
    ("deu", "tess_lang_german"),
    ("por", "tess_lang_portuguese"),
    ("ita", "tess_lang_italian"),
    ("jpn", "tess_lang_japanese"),
    ("kor", "tess_lang_korean"),
    ("ara", "tess_lang_arabic"),
    ("rus", "tess_lang_russian"),
)


@dataclass(frozen=True, slots=True)
class OllamaModelPreset:
    model: str
    label: str
    summary: str
    recommended: bool = False


@dataclass(frozen=True, slots=True)
class TesseractWordBox:
    block_num: int
    par_num: int
    line_num: int
    word_num: int
    left: int
    top: int
    width: int
    height: int
    text: str


@dataclass(slots=True)
class TesseractLineLayout:
    block_num: int
    par_num: int
    line_num: int
    left: int = 0
    top: int = 0
    width: int = 0
    height: int = 0
    words: list[TesseractWordBox] = field(default_factory=list)


@dataclass(frozen=True, slots=True)
class OcrTextPage:
    page_number: int
    image_path: Path
    text: str
    preserve_image: bool = False


@dataclass(frozen=True, slots=True)
class StructuredOcrBlock:
    kind: str
    text: str = ""
    image_path: Path | None = None


OLLAMA_MODEL_PRESETS = (
    OllamaModelPreset(
        model="glm-ocr",
        label="GLM-OCR",
        summary="Recommended: 2.2GB, 128K context, dedicated OCR model for complex documents with very low latency.",
        recommended=True,
    ),
    OllamaModelPreset(
        model="qwen3-vl:4b",
        label="Qwen3-VL 4B",
        summary="3.3GB, 256K context, strong multilingual OCR and document understanding with much lower latency than 7B models.",
    ),
    OllamaModelPreset(
        model="qwen3-vl:2b",
        label="Qwen3-VL 2B",
        summary="1.9GB, 256K context, fastest general vision fallback when you want lower latency over maximum accuracy.",
    ),
    OllamaModelPreset(
        model="openbmb/minicpm-v4.5:8b",
        label="MiniCPM-V 4.5 8B",
        summary="4.8GB model blob + 1.1GB projector (5.9GB total), built on Qwen3-8B, supports up to 1.8M-pixel images, strong OCRBench and OmniDocBench document parsing.",
    ),
    OllamaModelPreset(
        model="qwen2.5vl:3b",
        label="Qwen2.5-VL 3B",
        summary="3.2GB, document-friendly edge model with good layout understanding and better latency than the 7B variant.",
    ),
    OllamaModelPreset(
        model="qwen2.5vl:7b",
        label="Qwen2.5-VL 7B",
        summary="6.0GB, strong document/layout accuracy, but notably slower for page-by-page OCR than the lighter presets above.",
    ),
    OllamaModelPreset(
        model="qwen3-vl:8b",
        label="Qwen3-VL 8B",
        summary="6.1GB, 256K context, higher-capacity general VLM with improved OCR across 32 languages, but slower than GLM-OCR or Qwen3-VL 4B.",
    ),
    OllamaModelPreset(
        model="deepseek-ocr:3b",
        label="DeepSeek-OCR 3B",
        summary="6.7GB, OCR-specific model with token-efficient extraction and markdown-oriented prompts, but heavier than GLM-OCR.",
    ),
    OllamaModelPreset(
        model="richardyoung/olmocr2:7b-q8",
        label="olmOCR-2 7B 1025 Q8",
        summary="9.5GB, Qwen2.5-VL-7B OCR fine-tune scoring 82.4 on olmOCR-Bench; accurate but much heavier than GLM-OCR.",
    ),
)


def normalize_ollama_url(base_url: str) -> str:
    value = (base_url or DEFAULT_OLLAMA_URL).strip()
    return value.rstrip("/") or DEFAULT_OLLAMA_URL


def ollama_install_command() -> str:
    return OLLAMA_WINDOWS_INSTALL_COMMAND


def ollama_run_command() -> str:
    return OLLAMA_WINDOWS_RUN_COMMAND


def tesseract_install_command() -> str:
    return TESSERACT_WINDOWS_INSTALL_COMMAND


def ocrmypdf_install_command() -> str:
    return OCRMYPDF_WINDOWS_INSTALL_COMMAND


def surya_install_command() -> str:
    return SURYA_INSTALL_COMMAND


def _app_data_dir() -> Path:
    local_app_data = os.getenv("LOCALAPPDATA")
    if local_app_data:
        return Path(local_app_data) / "UniversalScreenshotTaker"
    return Path.home() / ".universal-screenshot-taker"


def tesseract_managed_dir() -> Path:
    return _app_data_dir() / "tesseract"


def tesseract_managed_tessdata_dir() -> Path:
    return tesseract_managed_dir() / "downloads" / "tessdata"


def tesseract_runtime_root() -> Path:
    return tesseract_managed_dir() / "runtime"


def tesseract_runtime_tessdata_dir() -> Path:
    return tesseract_runtime_root() / "tessdata"


def find_tesseract_executable() -> str | None:
    executable = shutil.which("tesseract")
    if executable:
        return executable
    candidates = (
        Path("C:/Program Files/Tesseract-OCR/tesseract.exe"),
        Path("C:/Program Files (x86)/Tesseract-OCR/tesseract.exe"),
    )
    for candidate in candidates:
        if candidate.exists():
            return str(candidate)
    return None


def find_ghostscript_executable() -> str | None:
    executable = shutil.which("gswin64c") or shutil.which("gswin32c")
    if executable:
        return executable
    search_roots = (
        Path("C:/Program Files/gs"),
        Path("C:/Program Files (x86)/gs"),
    )
    patterns = ("gs*/bin/gswin64c.exe", "gs*/bin/gswin32c.exe")
    for root in search_roots:
        if not root.exists():
            continue
        for pattern in patterns:
            matches = sorted(root.glob(pattern), reverse=True)
            if matches:
                return str(matches[0])
    return None


def ocrmypdf_available() -> bool:
    return importlib.util.find_spec("ocrmypdf") is not None


def surya_available() -> bool:
    if importlib.util.find_spec("surya") is not None:
        pass
    else:
        executable_name = "surya_ocr.exe" if os.name == "nt" else "surya_ocr"
        candidate = Path(sys.executable).resolve().parent / executable_name
        if not candidate.exists() and shutil.which("surya_ocr") is None:
            return False
    try:
        transformers_version = importlib.metadata.version("transformers")
    except importlib.metadata.PackageNotFoundError:
        return False
    major_match = re.match(r"^\s*(\d+)", transformers_version)
    if major_match and int(major_match.group(1)) >= 5:
        return False
    return True


def tesseract_system_tessdata_dir() -> Path | None:
    env_prefix = os.getenv("TESSDATA_PREFIX", "").strip()
    if env_prefix:
        prefix_path = Path(env_prefix)
        candidates = (prefix_path, prefix_path / "tessdata")
        for candidate in candidates:
            if candidate.exists() and candidate.is_dir():
                return candidate

    tesseract = find_tesseract_executable()
    if not tesseract:
        return None
    candidate = Path(tesseract).resolve().parent / "tessdata"
    if candidate.exists() and candidate.is_dir():
        return candidate
    return None


def _traineddata_codes(directory: Path | None) -> list[str]:
    if directory is None or not directory.exists():
        return []
    values: list[str] = []
    seen: set[str] = set()
    for path in sorted(directory.glob("*.traineddata")):
        code = path.stem.strip()
        if code and code not in seen:
            seen.add(code)
            values.append(code)
    return values


def build_tesseract_runtime_tessdata() -> Path | None:
    source_dirs: list[Path] = []
    system_dir = tesseract_system_tessdata_dir()
    if system_dir is not None:
        source_dirs.append(system_dir)
    managed_dir = tesseract_managed_tessdata_dir()
    if managed_dir.exists():
        source_dirs.append(managed_dir)
    if not source_dirs:
        return None

    runtime_root = tesseract_runtime_root()
    runtime_tessdata = tesseract_runtime_tessdata_dir()
    if runtime_root.exists():
        shutil.rmtree(runtime_root, ignore_errors=True)
    ensure_directory(runtime_tessdata)

    for source_dir in source_dirs:
        for path in sorted(source_dir.rglob("*")):
            relative_path = path.relative_to(source_dir)
            target_path = runtime_tessdata / relative_path
            if path.is_dir():
                ensure_directory(target_path)
            elif path.is_file():
                ensure_directory(target_path.parent)
                shutil.copy2(path, target_path)
    return runtime_tessdata


def tesseract_environment() -> dict[str, str]:
    env = os.environ.copy()
    runtime_tessdata = build_tesseract_runtime_tessdata()
    if runtime_tessdata is not None:
        env["TESSDATA_PREFIX"] = str(runtime_tessdata)
    return env


def list_tesseract_languages() -> list[str]:
    languages: list[str] = []
    seen: set[str] = set()
    for directory in (tesseract_system_tessdata_dir(), tesseract_managed_tessdata_dir()):
        for code in _traineddata_codes(directory):
            if code not in seen:
                seen.add(code)
                languages.append(code)
    return languages


def _parse_tesseract_language_codes(languages: str) -> list[str]:
    values: list[str] = []
    seen: set[str] = set()
    for part in languages.replace(",", "+").split("+"):
        code = part.strip()
        if code and code not in seen:
            seen.add(code)
            values.append(code)
    return values


def resolve_tesseract_languages(
    languages: str,
    available_languages: Iterable[str] | None = None,
) -> tuple[str, list[str]]:
    requested_codes = _parse_tesseract_language_codes(languages) or ["eng"]
    available = [code.strip() for code in (available_languages or []) if code and str(code).strip()]
    if not available:
        return "+".join(requested_codes), []

    valid_codes = [code for code in requested_codes if code in available]
    missing_codes = [code for code in requested_codes if code not in available]
    if valid_codes:
        return "+".join(valid_codes), missing_codes

    fallback = "eng" if "eng" in available else next((code for code in available if code != "osd"), available[0])
    return fallback, missing_codes or requested_codes


def install_tesseract_languages(
    languages: Iterable[str],
    progress_callback: ProgressCallback | None = None,
) -> list[Path]:
    requested_codes = [code for code in languages if str(code).strip()]
    if not requested_codes:
        return []

    managed_tessdata = ensure_directory(tesseract_managed_tessdata_dir())
    installed_paths: list[Path] = []
    total = len(requested_codes)
    for index, code in enumerate(requested_codes, start=1):
        normalized_code = str(code).strip()
        if progress_callback is not None:
            progress_callback(index - 1, total, f"Downloading {normalized_code}.traineddata...")
        url = f"{TESSERACT_TESSDATA_FAST_BASE_URL}/{normalized_code}.traineddata"
        target_path = managed_tessdata / f"{normalized_code}.traineddata"
        temp_path = target_path.with_suffix(".tmp")
        try:
            with urllib.request.urlopen(url, timeout=300.0) as response, temp_path.open("wb") as handle:
                shutil.copyfileobj(response, handle)
        except urllib.error.HTTPError as exc:
            if temp_path.exists():
                temp_path.unlink(missing_ok=True)
            if exc.code == 404:
                raise RuntimeError(
                    f"Tesseract language '{normalized_code}' is not available in the official tessdata_fast repository."
                ) from exc
            details = exc.read().decode("utf-8", errors="replace").strip()
            raise RuntimeError(details or f"Could not download Tesseract language '{normalized_code}'.") from exc
        except urllib.error.URLError as exc:
            if temp_path.exists():
                temp_path.unlink(missing_ok=True)
            raise RuntimeError(str(exc.reason) or f"Could not download Tesseract language '{normalized_code}'.") from exc
        temp_path.replace(target_path)
        installed_paths.append(target_path)
        if progress_callback is not None:
            progress_callback(index, total, f"Installed {normalized_code}.traineddata")

    build_tesseract_runtime_tessdata()
    return installed_paths


def _request_json(
    method: str,
    url: str,
    payload: dict[str, object] | None = None,
    timeout: float = 30.0,
) -> dict[str, object]:
    data = None if payload is None else json.dumps(payload).encode("utf-8")
    request = urllib.request.Request(
        url,
        data=data,
        headers={"Content-Type": "application/json"},
        method=method,
    )
    try:
        with urllib.request.urlopen(request, timeout=timeout) as response:
            return json.loads(response.read().decode("utf-8"))
    except urllib.error.HTTPError as exc:
        details = exc.read().decode("utf-8", errors="replace").strip()
        raise RuntimeError(details or f"Ollama request failed with HTTP {exc.code}.") from exc
    except urllib.error.URLError as exc:
        raise RuntimeError(str(exc.reason) or "Could not connect to Ollama.") from exc


def _request_json_stream(
    method: str,
    url: str,
    payload: dict[str, object] | None = None,
    timeout: float = 30.0,
) -> Iterable[dict[str, object]]:
    data = None if payload is None else json.dumps(payload).encode("utf-8")
    request = urllib.request.Request(
        url,
        data=data,
        headers={"Content-Type": "application/json"},
        method=method,
    )
    try:
        with urllib.request.urlopen(request, timeout=timeout) as response:
            for raw_line in response:
                line = raw_line.decode("utf-8", errors="replace").strip()
                if not line:
                    continue
                yield json.loads(line)
    except urllib.error.HTTPError as exc:
        details = exc.read().decode("utf-8", errors="replace").strip()
        raise RuntimeError(details or f"Ollama request failed with HTTP {exc.code}.") from exc
    except urllib.error.URLError as exc:
        raise RuntimeError(str(exc.reason) or "Could not connect to Ollama.") from exc


def ollama_version(base_url: str) -> str:
    payload = _request_json("GET", f"{normalize_ollama_url(base_url)}/api/version", timeout=5.0)
    version = str(payload.get("version", "")).strip()
    return version or "unknown"


def list_ollama_models(base_url: str) -> list[dict[str, object]]:
    payload = _request_json("GET", f"{normalize_ollama_url(base_url)}/api/tags", timeout=10.0)
    models = payload.get("models", [])
    if isinstance(models, list):
        return [entry for entry in models if isinstance(entry, dict)]
    return []


def ollama_model_installed(base_url: str, model_name: str) -> bool:
    target = model_name.strip()
    for entry in list_ollama_models(base_url):
        if target in {str(entry.get("name", "")).strip(), str(entry.get("model", "")).strip()}:
            return True
    return False


def pull_ollama_model(base_url: str, model_name: str, progress_callback: ProgressCallback | None = None) -> None:
    request = urllib.request.Request(
        f"{normalize_ollama_url(base_url)}/api/pull",
        data=json.dumps({"model": model_name, "stream": True}).encode("utf-8"),
        headers={"Content-Type": "application/json"},
        method="POST",
    )
    try:
        with urllib.request.urlopen(request, timeout=3600) as response:
            for raw_line in response:
                line = raw_line.decode("utf-8", errors="replace").strip()
                if not line:
                    continue
                event = json.loads(line)
                status = str(event.get("status", "")).strip() or f"Pulling {model_name}"
                total = int(event.get("total", 0) or 0)
                completed = int(event.get("completed", 0) or 0)
                if progress_callback is not None:
                    progress_callback(completed, total, status)
    except urllib.error.HTTPError as exc:
        details = exc.read().decode("utf-8", errors="replace").strip()
        raise RuntimeError(details or f"Could not pull model {model_name}.") from exc
    except urllib.error.URLError as exc:
        raise RuntimeError(str(exc.reason) or "Could not connect to Ollama.") from exc


def test_ollama_connection(base_url: str, model_name: str) -> str:
    version = ollama_version(base_url)
    if not ollama_model_installed(base_url, model_name):
        raise RuntimeError(f"Ollama {version} is reachable, but model {model_name} is not installed yet.")
    response = _request_json(
        "POST",
        f"{normalize_ollama_url(base_url)}/api/chat",
        payload={
            "model": model_name,
            "stream": False,
            "think": False,
            "messages": [
                {
                    "role": "user",
                    "content": "Reply with OK only.",
                }
            ],
            "options": {"temperature": 0},
        },
        timeout=180.0,
    )
    content = str(((response.get("message") or {}) if isinstance(response.get("message"), dict) else {}).get("content", "")).strip()
    if not content:
        raise RuntimeError(f"Ollama {version} responded, but the model returned an empty reply.")
    return version


def export_tesseract_ocr_pdf(
    image_paths: Iterable[Path],
    pdf_path: Path,
    languages: str = "eng",
    progress_callback: ProgressCallback | None = None,
) -> Path:
    ordered_paths = [Path(path) for path in image_paths]
    if not ordered_paths:
        raise ValueError("No screenshots are available to export.")

    tesseract = find_tesseract_executable()
    if not tesseract:
        raise RuntimeError(
            "Tesseract OCR is not installed or not in PATH. "
            f"Install it and try again. Suggested command: {tesseract_install_command()}"
        )

    ensure_directory(pdf_path.parent)
    available_languages = list_tesseract_languages()
    language_value, missing_languages = resolve_tesseract_languages(languages, available_languages)
    if not language_value:
        raise RuntimeError(
            "No valid Tesseract languages are available. "
            f"Requested: {languages or 'eng'}. Available: {', '.join(available_languages) or 'none'}."
        )
    if missing_languages and not available_languages:
        raise RuntimeError(
            "Tesseract language data could not be enumerated. "
            f"Requested: {languages or 'eng'}. Install the language pack and try again."
        )
    tesseract_env = tesseract_environment()

    with tempfile.TemporaryDirectory(prefix="ust_tesseract_") as temp_dir_name:
        temp_dir = Path(temp_dir_name)
        page_pdfs: list[Path] = []
        total = len(ordered_paths)
        for index, image_path in enumerate(ordered_paths, start=1):
            if progress_callback is not None:
                progress_callback(index - 1, total, f"OCR {index}/{total}: {image_path.name}")
            output_stem = temp_dir / f"page_{index:04d}"
            result = subprocess.run(
                [
                    tesseract,
                    str(image_path),
                    str(output_stem),
                    "-l",
                    language_value,
                    "--oem",
                    "1",
                    "--dpi",
                    "300",
                    "pdf",
                ],
                capture_output=True,
                text=False,
                check=False,
                env=tesseract_env,
            )
            if result.returncode != 0:
                error_text = (_decode_subprocess_output(result.stderr) or _decode_subprocess_output(result.stdout)).strip()
                raise RuntimeError(error_text or f"Tesseract failed while processing {image_path.name}.")
            page_pdfs.append(output_stem.with_suffix(".pdf"))

        writer = PdfWriter()
        for page_pdf in page_pdfs:
            reader = PdfReader(str(page_pdf))
            for page in reader.pages:
                writer.add_page(page)
        with pdf_path.open("wb") as handle:
            writer.write(handle)

    if not pdf_has_extractable_text(pdf_path):
        raise RuntimeError("Tesseract finished, but the exported PDF does not contain selectable text.")

    if progress_callback is not None:
        progress_callback(len(ordered_paths), len(ordered_paths), "Tesseract OCR complete. Searchable text verified.")
    return pdf_path


def _image_bytes_base64(image_path: Path) -> str:
    import base64

    return base64.b64encode(image_path.read_bytes()).decode("ascii")


def _decode_subprocess_output(value: bytes | str | None) -> str:
    if value is None:
        return ""
    if isinstance(value, str):
        decoded = value
    else:
        decoded = ""
        for encoding in ("utf-8", locale.getpreferredencoding(False) or "utf-8", "cp1252"):
            try:
                decoded = value.decode(encoding)
                break
            except UnicodeDecodeError:
                continue
        if not decoded:
            decoded = value.decode("utf-8", errors="backslashreplace")
    decoded = decoded.replace("\ufffd", "?")
    sanitized_chars: list[str] = []
    for char in decoded:
        if char in "\r\n\t" or ord(char) >= 32:
            sanitized_chars.append(char)
        else:
            sanitized_chars.append("?")
    return "".join(sanitized_chars)


def _strip_html_artifacts(value: str) -> str:
    if value is None:
        return ""
    cleaned = str(value)
    cleaned = re.sub(r"(?is)^\s*```(?:html|xml|markdown|text)?\s*", "", cleaned)
    cleaned = re.sub(r"(?is)\s*```\s*$", "", cleaned)
    if "<" in cleaned and ">" in cleaned:
        cleaned = re.sub(
            r"(?is)<\s*/?\s*(br|p|div|li|tr|h[1-6]|table|tbody|thead|ul|ol)\b[^>]*>",
            "\n",
            cleaned,
        )
        cleaned = re.sub(r"(?is)<\s*/?\s*(html|body|span|strong|b|i|em|section|article|main|header|footer)\b[^>]*>", " ", cleaned)
        cleaned = re.sub(r"(?is)<[^>]+>", "", cleaned)
    cleaned = html.unescape(cleaned)
    cleaned = cleaned.replace("\xa0", " ")
    return cleaned


def _strip_ollama_reasoning_artifacts(value: str) -> str:
    if value is None:
        return ""
    cleaned = str(value)
    cleaned = re.sub(r"(?is)<\s*think(?:ing)?\b[^>]*>.*?</\s*think(?:ing)?\s*>", "", cleaned)
    cleaned = re.sub(r"(?is)^\s*```(?:thinking|reasoning)\s*.*?```\s*", "", cleaned)
    cleaned = re.sub(r"(?im)^\s*(?:chain[- ]of[- ]thought|reasoning|analysis)\s*:\s*$", "", cleaned)
    return cleaned.strip()


def _normalize_ocr_text(value: str) -> str:
    if value is None:
        return ""
    cleaned = _strip_html_artifacts(_strip_ollama_reasoning_artifacts(value))
    lines = [line.rstrip() for line in cleaned.replace("\r\n", "\n").replace("\r", "\n").split("\n")]
    while lines and not lines[0].strip():
        lines.pop(0)
    while lines and not lines[-1].strip():
        lines.pop()
    return "\n".join(lines)


def _coerce_ollama_content(value: object) -> str:
    if value is None:
        return ""
    if isinstance(value, str):
        return value
    if isinstance(value, list):
        parts: list[str] = []
        for entry in value:
            if isinstance(entry, dict):
                text_value = entry.get("text")
                if isinstance(text_value, str):
                    parts.append(text_value)
            elif isinstance(entry, str):
                parts.append(entry)
        return "".join(parts)
    if isinstance(value, dict):
        text_value = value.get("text")
        if isinstance(text_value, str):
            return text_value
    return str(value)


def pdf_has_extractable_text(pdf_path: Path) -> bool:
    reader = PdfReader(str(pdf_path))
    for page in reader.pages:
        extracted_text = (page.extract_text() or "").strip()
        if extracted_text:
            return True
    return False


def _ollama_ocr_prompt(model_name: str, prompt: str) -> str:
    if prompt != DEFAULT_OCR_PROMPT:
        return prompt
    normalized_model = model_name.strip().lower()
    if normalized_model.startswith("glm-ocr"):
        return (
            "Text Recognition. Return only the main body text as plain text. Preserve real paragraph breaks. "
            "If the page contains a chart, diagram, table, illustration, equation block, or other important visual that should be preserved in rich exports, "
            "add a separate line exactly [[KEEP_PAGE_IMAGE]]. "
            "If the page starts or ends mid-paragraph, continue naturally without inventing a new paragraph. "
            "Ignore repeated headers, footers, page numbers, and decorative page elements."
        )
    if normalized_model.startswith("deepseek-ocr"):
        return (
            "Extract only the main body text in the image as plain text. Preserve real paragraph breaks. "
            "If the page contains a chart, diagram, table, illustration, equation block, or other important visual that should be preserved in rich exports, "
            "add a separate line exactly [[KEEP_PAGE_IMAGE]]. "
            "If the page starts or ends mid-paragraph, continue naturally without inventing a new paragraph. "
            "Ignore repeated headers, footers, page numbers, and decorative page elements."
        )
    if "olmocr2" in normalized_model:
        return (
            "Extract the main body text from this image, preserving reading order and paragraph structure. "
            "If the page contains a chart, diagram, table, illustration, equation block, or other important visual that should be preserved in rich exports, "
            "add a separate line exactly [[KEEP_PAGE_IMAGE]]. "
            "If the page starts or ends mid-paragraph, continue naturally without inventing a new paragraph. "
            "Return plain text only. Ignore repeated headers, footers, page numbers, and decorative page elements."
        )
    return DEFAULT_OCR_PROMPT


def extract_text_with_ollama(
    base_url: str,
    model_name: str,
    image_path: Path,
    prompt: str = DEFAULT_OCR_PROMPT,
    status_callback: StatusCallback | None = None,
) -> str:
    effective_prompt = _ollama_ocr_prompt(model_name, prompt)
    if status_callback is not None:
        status_callback(f"Running {model_name} on {image_path.name}...")

    content_parts: list[str] = []
    received_chars = 0
    last_status_update = 0.0
    for event in _request_json_stream(
        "POST",
        f"{normalize_ollama_url(base_url)}/api/chat",
        payload={
            "model": model_name,
            "stream": True,
            "think": False,
            "messages": [
                {
                    "role": "user",
                    "content": effective_prompt,
                    "images": [_image_bytes_base64(image_path)],
                }
            ],
            "options": {"temperature": 0},
            "keep_alive": "10m",
        },
        timeout=900.0,
    ):
        message = event.get("message")
        chunk = _coerce_ollama_content(message.get("content") if isinstance(message, dict) else "")
        if chunk:
            content_parts.append(chunk)
            received_chars += len(chunk)
            if status_callback is not None:
                now = time.monotonic()
                if last_status_update == 0.0:
                    status_callback(f"Receiving OCR text for {image_path.name}...")
                    last_status_update = now
                elif now - last_status_update >= 1.0:
                    status_callback(f"Receiving OCR text for {image_path.name} ({received_chars} chars)...")
                    last_status_update = now
        if bool(event.get("done")):
            break

    if status_callback is not None:
        status_callback(f"Finalizing OCR for {image_path.name}...")

    content = "".join(content_parts).strip()
    normalized = _normalize_ocr_text(content)
    if not normalized:
        raise RuntimeError(f"Ollama returned empty OCR text for {image_path.name}.")
    return normalized


def _wrap_overlay_lines(text: str, max_chars: int) -> list[str]:
    normalized = _normalize_ocr_text(text)
    if not normalized:
        return [""]
    max_chars = max(12, max_chars)
    wrapped_lines: list[str] = []
    for raw_line in normalized.splitlines():
        line = raw_line.rstrip()
        if not line.strip():
            wrapped_lines.append("")
            continue
        if len(line) <= max_chars:
            wrapped_lines.append(line)
            continue
        indent = len(line) - len(line.lstrip(" "))
        indent_text = " " * indent
        wrapped_lines.extend(
            textwrap.wrap(
                line,
                width=max_chars,
                break_long_words=False,
                replace_whitespace=False,
                drop_whitespace=False,
                subsequent_indent=indent_text,
            )
            or [line]
        )
    return wrapped_lines


def _tesseract_tsv_rows(image_path: Path) -> list[list[str]]:
    tesseract = find_tesseract_executable()
    if not tesseract:
        return []

    available_languages = list_tesseract_languages()
    language_candidates = [code for code in available_languages if code != "osd"]
    language_value = "+".join(language_candidates) if language_candidates else "eng"
    result = subprocess.run(
        [
            tesseract,
            str(image_path),
            "stdout",
            "-l",
            language_value,
            "--psm",
            "6",
            "tsv",
        ],
        capture_output=True,
        text=False,
        check=False,
        env=tesseract_environment(),
    )
    if result.returncode != 0:
        return []

    stdout_text = _decode_subprocess_output(result.stdout)
    lines = [line for line in stdout_text.splitlines() if line.strip()]
    if not lines:
        return []
    return [raw_line.split("\t") for raw_line in lines[1:]]


def _line_layouts_from_tesseract_rows(rows: Iterable[list[str]]) -> list[TesseractLineLayout]:
    layouts_by_key: dict[tuple[int, int, int], TesseractLineLayout] = {}
    ordered_keys: list[tuple[int, int, int]] = []
    for columns in rows:
        if len(columns) < 12:
            continue
        try:
            level = int(columns[0].strip())
            block_num = int(columns[2])
            par_num = int(columns[3])
            line_num = int(columns[4])
            left = int(columns[6])
            top = int(columns[7])
            width = int(columns[8])
            height = int(columns[9])
        except ValueError:
            continue
        key = (block_num, par_num, line_num)
        layout = layouts_by_key.get(key)
        if layout is None:
            layout = TesseractLineLayout(block_num=block_num, par_num=par_num, line_num=line_num)
            layouts_by_key[key] = layout
            ordered_keys.append(key)
        if level == 4:
            layout.left = left
            layout.top = top
            layout.width = width
            layout.height = height
            continue
        if level != 5 or width <= 0 or height <= 0:
            continue
        text_value = columns[11].strip()
        if not text_value:
            continue
        layout.words.append(
            TesseractWordBox(
                block_num=block_num,
                par_num=par_num,
                line_num=line_num,
                word_num=int(columns[5]) if str(columns[5]).strip().isdigit() else len(layout.words) + 1,
                left=left,
                top=top,
                width=width,
                height=height,
                text=text_value,
            )
        )

    layouts: list[TesseractLineLayout] = []
    for key in ordered_keys:
        layout = layouts_by_key[key]
        if layout.words and (layout.width <= 0 or layout.height <= 0):
            left = min(word.left for word in layout.words)
            top = min(word.top for word in layout.words)
            right = max(word.left + word.width for word in layout.words)
            bottom = max(word.top + word.height for word in layout.words)
            layout.left = left
            layout.top = top
            layout.width = max(1, right - left)
            layout.height = max(1, bottom - top)
        layouts.append(layout)
    return layouts


def _word_boxes_from_tesseract_rows(rows: Iterable[list[str]]) -> list[TesseractWordBox]:
    boxes: list[TesseractWordBox] = []
    seen_boxes: set[tuple[int, int, int, int, int, int, int, int]] = set()
    for layout in _line_layouts_from_tesseract_rows(rows):
        for word in layout.words:
            key = (
                word.block_num,
                word.par_num,
                word.line_num,
                word.word_num,
                word.left,
                word.top,
                word.width,
                word.height,
            )
            if key in seen_boxes:
                continue
            seen_boxes.add(key)
            boxes.append(word)
    return boxes


def _line_level_boxes_from_tesseract(image_path: Path) -> list[tuple[int, int, int, int]]:
    rows = _tesseract_tsv_rows(image_path)
    boxes: list[tuple[int, int, int, int]] = []
    seen_boxes: set[tuple[int, int, int, int]] = set()
    for layout in _line_layouts_from_tesseract_rows(rows):
        box = (layout.left, layout.top, layout.width, layout.height)
        if layout.width <= 0 or layout.height <= 0 or box in seen_boxes:
            continue
        seen_boxes.add(box)
        boxes.append(box)
    return boxes


def _normalized_visible_lines(text: str) -> list[str]:
    values: list[str] = []
    for raw_line in _normalize_ocr_text(text).splitlines():
        line = raw_line.rstrip()
        if line.strip():
            values.append(line)
    return values


def _normalized_visible_tokens(text: str) -> list[str]:
    tokens: list[str] = []
    for token in re.findall(r"\S+", _normalize_ocr_text(text)):
        cleaned_token = token.strip()
        if not cleaned_token:
            continue
        if re.fullmatch(r"(?is)</?[^>]+>", cleaned_token):
            continue
        tokens.append(cleaned_token)
    return tokens


def _normalize_alignment_token(value: str) -> str:
    normalized = re.sub(r"[\W_]+", "", (value or "").casefold(), flags=re.UNICODE)
    return normalized or (value or "").strip().casefold()


def _assign_tokens_to_word_boxes(text: str, word_boxes: list[TesseractWordBox]) -> list[tuple[TesseractWordBox, str]]:
    tokens = _normalized_visible_tokens(text)
    if not tokens or not word_boxes:
        return []

    box_buckets: list[list[str]] = [[] for _ in word_boxes]
    box_tokens = [_normalize_alignment_token(box.text) for box in word_boxes]
    ai_tokens = [_normalize_alignment_token(token) for token in tokens]
    matcher = SequenceMatcher(a=box_tokens, b=ai_tokens, autojunk=False)

    for tag, box_start, box_end, token_start, token_end in matcher.get_opcodes():
        if tag == "equal":
            pair_count = min(box_end - box_start, token_end - token_start)
            for offset in range(pair_count):
                box_buckets[box_start + offset].append(tokens[token_start + offset])
            continue
        if token_start >= token_end:
            continue
        if box_start >= box_end:
            anchor = min(max(box_start - 1, 0), len(word_boxes) - 1)
            box_buckets[anchor].extend(tokens[token_start:token_end])
            continue
        box_span = box_end - box_start
        token_span = token_end - token_start
        for offset, token in enumerate(tokens[token_start:token_end]):
            target_offset = min(box_span - 1, int((offset * box_span) / max(token_span, 1)))
            box_buckets[box_start + target_offset].append(token)

    assignments: list[tuple[TesseractWordBox, str]] = []
    for box, bucket in zip(word_boxes, box_buckets, strict=False):
        if not bucket:
            continue
        assignments.append((box, " ".join(bucket)))
    return assignments


def _layout_visible_text(layout: TesseractLineLayout) -> str:
    return " ".join(word.text.strip() for word in layout.words if word.text.strip())


def _assign_lines_to_layouts(text: str, line_layouts: list[TesseractLineLayout]) -> list[tuple[TesseractLineLayout, str]]:
    visible_lines = _normalized_visible_lines(text)
    usable_layouts = [layout for layout in line_layouts if layout.words]
    if not visible_lines or not usable_layouts:
        return []

    line_buckets: list[list[str]] = [[] for _ in usable_layouts]
    tesseract_lines = [_normalize_alignment_token(_layout_visible_text(layout)) for layout in usable_layouts]
    ai_lines = [_normalize_alignment_token(line) for line in visible_lines]
    matcher = SequenceMatcher(a=tesseract_lines, b=ai_lines, autojunk=False)

    for tag, layout_start, layout_end, line_start, line_end in matcher.get_opcodes():
        if tag == "equal":
            pair_count = min(layout_end - layout_start, line_end - line_start)
            for offset in range(pair_count):
                line_buckets[layout_start + offset].append(visible_lines[line_start + offset])
            continue
        if line_start >= line_end:
            continue
        if layout_start >= layout_end:
            anchor = min(max(layout_start - 1, 0), len(usable_layouts) - 1)
            line_buckets[anchor].extend(visible_lines[line_start:line_end])
            continue
        layout_span = layout_end - layout_start
        line_span = line_end - line_start
        for offset, line_text in enumerate(visible_lines[line_start:line_end]):
            target_offset = min(layout_span - 1, int((offset * layout_span) / max(line_span, 1)))
            line_buckets[layout_start + target_offset].append(line_text)

    assignments: list[tuple[TesseractLineLayout, str]] = []
    for layout, bucket in zip(usable_layouts, line_buckets, strict=False):
        merged_line = " ".join(part.strip() for part in bucket if part.strip()).strip()
        if merged_line:
            assignments.append((layout, merged_line))
    return assignments


def _write_positioned_word_overlay(
    pdf: canvas.Canvas,
    text: str,
    page_height: float,
    word_boxes: list[TesseractWordBox],
) -> bool:
    assignments = _assign_tokens_to_word_boxes(text, word_boxes)
    if not assignments:
        return False

    for box, token_text in assignments:
        font_size = max(4.0, min(float(box.height) * 0.82, 42.0))
        baseline_y = max(page_height - float(box.top) - float(box.height) + (font_size * 0.92), font_size)
        text_object = pdf.beginText(float(box.left), baseline_y)
        text_object.setFont("Courier", font_size)
        if hasattr(text_object, "setTextRenderMode"):
            text_object.setTextRenderMode(3)
        natural_width = pdf.stringWidth(token_text, "Courier", font_size)
        if natural_width > 0:
            scale = max(20.0, min((float(box.width) / natural_width) * 100.0, 260.0))
            text_object.setHorizScale(scale)
        text_object.textLine(token_text)
        pdf.drawText(text_object)
    return True


def _write_positioned_layout_overlay(
    pdf: canvas.Canvas,
    text: str,
    page_width: float,
    page_height: float,
    line_layouts: list[TesseractLineLayout],
) -> tuple[str, int]:
    assignments = _assign_lines_to_layouts(text, line_layouts)
    if not assignments:
        return "fallback", 0

    aligned_word_boxes = 0
    aligned_line_boxes = 0
    for layout, line_text in assignments:
        if _write_positioned_word_overlay(pdf, line_text, page_height, layout.words):
            aligned_word_boxes += len(layout.words)
            continue
        if layout.width > 0 and layout.height > 0 and _write_positioned_line_overlay(
            pdf,
            line_text,
            page_width,
            page_height,
            [(layout.left, layout.top, layout.width, layout.height)],
        ):
            aligned_line_boxes += 1

    if aligned_word_boxes:
        return "word_boxes", aligned_word_boxes
    if aligned_line_boxes:
        return "line_boxes", aligned_line_boxes
    return "fallback", 0


def _write_positioned_line_overlay(
    pdf: canvas.Canvas,
    text: str,
    page_width: float,
    page_height: float,
    line_boxes: list[tuple[int, int, int, int]],
) -> bool:
    lines = _normalized_visible_lines(text)
    if not lines or not line_boxes:
        return False

    line_count = min(len(lines), len(line_boxes))
    for index in range(line_count):
        line_text = lines[index]
        left, top, width, height = line_boxes[index]
        font_size = max(4.0, min(float(height) * 0.82, 42.0))
        baseline_y = max(page_height - float(top) - float(height) + (font_size * 0.92), font_size)
        text_object = pdf.beginText(float(left), baseline_y)
        text_object.setFont("Courier", font_size)
        if hasattr(text_object, "setTextRenderMode"):
            text_object.setTextRenderMode(3)
        natural_width = pdf.stringWidth(line_text, "Courier", font_size)
        if natural_width > 0:
            scale = max(20.0, min((float(width) / natural_width) * 100.0, 220.0))
            text_object.setHorizScale(scale)
        text_object.textLine(line_text)
        pdf.drawText(text_object)
    return True


def _overlay_text_layout(text: str, page_width: float, page_height: float) -> tuple[list[str], float, float, float, float]:
    margin_x = max(18.0, page_width * 0.04)
    margin_y = max(18.0, page_height * 0.04)
    content_width = max(80.0, page_width - (margin_x * 2.0))
    content_height = max(80.0, page_height - (margin_y * 2.0))

    best_lines = [""]
    best_font_size = 6.0
    best_leading = 7.2
    font_size = 24.0
    while font_size >= 5.0:
        estimated_char_width = max(font_size * 0.52, 1.0)
        max_chars = max(12, int(content_width / estimated_char_width))
        lines = _wrap_overlay_lines(text, max_chars)
        leading = max(font_size * 1.18, font_size + 1.0)
        best_lines = lines
        best_font_size = font_size
        best_leading = leading
        if len(lines) * leading <= content_height:
            return lines, font_size, leading, margin_x, margin_y
        font_size -= 0.5

    line_count = max(len(best_lines), 1)
    compressed_leading = max(3.2, content_height / line_count)
    compressed_font_size = max(3.0, min(best_font_size, compressed_leading / 1.18))
    return best_lines, compressed_font_size, compressed_leading, margin_x, margin_y


def _write_overlay_text(pdf: canvas.Canvas, text: str, page_width: float, page_height: float) -> None:
    lines, font_size, leading, margin_x, margin_y = _overlay_text_layout(text, page_width, page_height)
    start_y = max(page_height - margin_y - font_size, font_size + margin_y)
    text_object = pdf.beginText(margin_x, start_y)
    text_object.setFont("Courier", font_size)
    text_object.setLeading(leading)
    if hasattr(text_object, "setTextRenderMode"):
        text_object.setTextRenderMode(3)
    for line in lines:
        visible_line = line.rstrip()
        if len(visible_line) >= 10:
            natural_width = pdf.stringWidth(visible_line, "Courier", font_size)
            if 0 < natural_width < (page_width - margin_x * 2.0) * 0.98:
                char_count = max(len(visible_line) - 1, 1)
                extra_space = ((page_width - margin_x * 2.0) - natural_width) / char_count
                text_object.setCharSpace(max(0.0, extra_space))
            else:
                text_object.setCharSpace(0.0)
        else:
            text_object.setCharSpace(0.0)
        text_object.textLine(line)
    pdf.drawText(text_object)


def _write_best_effort_overlay_text(
    pdf: canvas.Canvas,
    text: str,
    image_path: Path,
    page_width: float,
    page_height: float,
) -> tuple[str, int]:
    rows = _tesseract_tsv_rows(image_path)
    line_layouts = _line_layouts_from_tesseract_rows(rows)
    overlay_mode, overlay_count = _write_positioned_layout_overlay(pdf, text, page_width, page_height, line_layouts)
    if overlay_mode != "fallback":
        return overlay_mode, overlay_count
    line_boxes = [(layout.left, layout.top, layout.width, layout.height) for layout in line_layouts if layout.width > 0 and layout.height > 0]
    if _write_positioned_line_overlay(pdf, text, page_width, page_height, line_boxes):
        return "line_boxes", len(line_boxes)
    return "fallback", 0


def export_ollama_ocr_pdf(
    image_paths: Iterable[Path],
    pdf_path: Path,
    base_url: str,
    model_name: str,
    progress_callback: ProgressCallback | None = None,
    prompt: str = DEFAULT_OCR_PROMPT,
) -> Path:
    ordered_paths = [Path(path) for path in image_paths]
    if not ordered_paths:
        raise ValueError("No screenshots are available to export.")
    if not ollama_model_installed(base_url, model_name):
        raise RuntimeError(
            f"Model {model_name} is not installed in Ollama. Pull it first from the app before starting OCR."
        )

    ensure_directory(pdf_path.parent)
    total = len(ordered_paths)
    pdf = canvas.Canvas(str(pdf_path))
    for index, image_path in enumerate(ordered_paths, start=1):
        if progress_callback is not None:
            progress_callback(index - 1, total, f"OCR {index}/{total}: {image_path.name}")
        status_callback = None
        if progress_callback is not None:
            status_callback = lambda status, completed=index - 1, total_pages=total: progress_callback(
                completed,
                total_pages,
                status,
            )
        text = extract_text_with_ollama(
            base_url,
            model_name,
            image_path,
            prompt=prompt,
            status_callback=status_callback,
        )
        if progress_callback is not None:
            progress_callback(index - 1, total, f"Writing OCR page {index}/{total}")
        with Image.open(image_path) as image:
            width, height = image.size
        page_width = float(width)
        page_height = float(height)
        pdf.setPageSize((page_width, page_height))
        pdf.drawImage(ImageReader(str(image_path)), 0, 0, width=page_width, height=page_height, mask="auto")
        overlay_mode, overlay_count = _write_best_effort_overlay_text(pdf, text, image_path, page_width, page_height)
        if progress_callback is not None:
            if overlay_mode == "word_boxes":
                progress_callback(index - 1, total, f"Aligned OCR text with {overlay_count} detected word boxes for {image_path.name}")
            elif overlay_mode == "line_boxes":
                progress_callback(index - 1, total, f"Aligned OCR text with {overlay_count} detected line boxes for {image_path.name}")
            else:
                progress_callback(index - 1, total, f"Tesseract could not determine OCR positions for {image_path.name}")
        if overlay_mode == "fallback":
            raise RuntimeError(
                f"Tesseract could not determine text positions for {image_path.name}. "
                "Install or repair Tesseract and retry so the OCR text can be placed correctly."
            )
        pdf.showPage()
        if progress_callback is not None:
            progress_callback(index, total, f"Added OCR page {index}/{total}")
    pdf.save()

    if progress_callback is not None:
        progress_callback(len(ordered_paths), len(ordered_paths), "Verifying searchable text layer...")
    if not pdf_has_extractable_text(pdf_path):
        raise RuntimeError("Ollama OCR finished, but the exported PDF does not contain selectable text.")

    if progress_callback is not None:
        progress_callback(len(ordered_paths), len(ordered_paths), "Ollama OCR complete. Searchable text verified.")
    return pdf_path


def _collect_ollama_ocr_pages(
    image_paths: Iterable[Path],
    base_url: str,
    model_name: str,
    progress_callback: ProgressCallback | None = None,
    prompt: str = DEFAULT_OCR_PROMPT,
) -> list[OcrTextPage]:
    ordered_paths = [Path(path) for path in image_paths]
    if not ordered_paths:
        raise ValueError("No screenshots are available to export.")
    if not ollama_model_installed(base_url, model_name):
        raise RuntimeError(
            f"Model {model_name} is not installed in Ollama. Pull it first from the app before starting OCR."
        )

    raw_pages: list[OcrTextPage] = []
    total = len(ordered_paths)
    for index, image_path in enumerate(ordered_paths, start=1):
        if progress_callback is not None:
            progress_callback(index - 1, total, f"OCR {index}/{total}: {image_path.name}")
        status_callback = None
        if progress_callback is not None:
            status_callback = lambda status, completed=index - 1, total_pages=total: progress_callback(
                completed,
                total_pages,
                status,
            )
        text = extract_text_with_ollama(
            base_url,
            model_name,
            image_path,
            prompt=prompt,
            status_callback=status_callback,
        )
        normalized_text = _normalize_ocr_text(text)
        if not normalized_text:
            raise RuntimeError(f"Ollama returned empty OCR text for {image_path.name}.")
        raw_pages.append(OcrTextPage(page_number=index, image_path=image_path, text=normalized_text))
        if progress_callback is not None:
            progress_callback(index, total, f"Collected OCR page {index}/{total}")
    return raw_pages


def _surya_command() -> list[str]:
    executable_name = "surya_ocr.exe" if os.name == "nt" else "surya_ocr"
    local_executable = Path(sys.executable).resolve().parent / executable_name
    if local_executable.exists():
        return [str(local_executable)]
    command = shutil.which("surya_ocr")
    if command:
        return [command]
    raise RuntimeError(
        "Surya OCR is not installed yet. "
        f"Install it with: {surya_install_command()}"
    )


def _surya_page_text(page_payload: dict[str, object]) -> str:
    text_lines = page_payload.get("text_lines")
    if isinstance(text_lines, list):
        lines: list[str] = []
        for entry in text_lines:
            if not isinstance(entry, dict):
                continue
            line_text = str(entry.get("text", "") or "").strip()
            if line_text:
                lines.append(line_text)
        if lines:
            return "\n".join(lines)
    return str(page_payload.get("text", "") or "").strip()


def _collect_surya_ocr_pages(
    image_paths: Iterable[Path],
    progress_callback: ProgressCallback | None = None,
) -> list[OcrTextPage]:
    ordered_paths = [Path(path) for path in image_paths]
    if not ordered_paths:
        raise ValueError("No screenshots are available to export.")
    command = _surya_command()

    pages: list[OcrTextPage] = []
    total = len(ordered_paths)
    with tempfile.TemporaryDirectory(prefix="ust_surya_") as temp_dir_name:
        temp_dir = Path(temp_dir_name)
        for index, image_path in enumerate(ordered_paths, start=1):
            output_dir = temp_dir / f"page_{index:04d}"
            ensure_directory(output_dir)
            if progress_callback is not None:
                progress_callback(index - 1, total, f"Running Surya OCR {index}/{total}: {image_path.name}")
            result = subprocess.run(
                [
                    *command,
                    str(image_path),
                    "--output_dir",
                    str(output_dir),
                ],
                capture_output=True,
                text=False,
                check=False,
                env=os.environ.copy(),
            )
            if result.returncode != 0:
                error_text = (_decode_subprocess_output(result.stderr) or _decode_subprocess_output(result.stdout)).strip()
                raise RuntimeError(error_text or f"Surya OCR failed while processing {image_path.name}.")

            result_candidates = sorted(output_dir.rglob("results.json"))
            results_path = result_candidates[0] if result_candidates else None
            if results_path is None or not results_path.exists():
                raise RuntimeError(f"Surya OCR did not produce results for {image_path.name}.")
            payload = json.loads(results_path.read_text(encoding="utf-8"))
            if not isinstance(payload, dict):
                raise RuntimeError(f"Surya OCR returned an invalid results file for {image_path.name}.")

            page_entries = payload.get(image_path.stem)
            if not isinstance(page_entries, list) or not page_entries:
                page_entries = next((value for value in payload.values() if isinstance(value, list) and value), None)
            if not isinstance(page_entries, list) or not page_entries:
                raise RuntimeError(f"Surya OCR returned no text for {image_path.name}.")

            page_payload = page_entries[0]
            if not isinstance(page_payload, dict):
                raise RuntimeError(f"Surya OCR returned malformed page data for {image_path.name}.")
            text = _normalize_ocr_text(_surya_page_text(page_payload))
            if not text:
                raise RuntimeError(f"Surya OCR returned empty text for {image_path.name}.")
            pages.append(OcrTextPage(page_number=index, image_path=image_path, text=text))
            if progress_callback is not None:
                progress_callback(index, total, f"Collected Surya OCR page {index}/{total}")
    return pages


def _ocr_text_document_title(target_path: Path) -> str:
    title = target_path.stem.strip()
    return title or "OCR Export"


def _normalize_export_text(value: str) -> str:
    normalized = _normalize_ocr_text(value).replace("\u00ad", "")
    if not normalized:
        return ""

    paragraphs = re.split(r"\n\s*\n", normalized)
    merged_paragraphs: list[str] = []
    for paragraph in paragraphs:
        lines = [line.strip() for line in paragraph.splitlines() if line.strip()]
        if not lines:
            continue
        merged = lines[0]
        for line in lines[1:]:
            if merged.endswith("-") and merged[:-1] and merged[:-1][-1].isalpha() and line and line[0].isalpha():
                merged = merged[:-1] + line
            else:
                merged = f"{merged} {line}"
        merged_paragraphs.append(_repair_intraword_hyphens(merged))
    return "\n\n".join(merged_paragraphs)


def _repair_intraword_hyphens(value: str) -> str:
    if "-" not in value:
        return value

    repaired_chars: list[str] = []
    length = len(value)
    for index, char in enumerate(value):
        if char != "-" or index == 0 or index >= length - 1:
            repaired_chars.append(char)
            continue

        prev_char = value[index - 1]
        next_char = value[index + 1]
        if not (prev_char.isalpha() and next_char.isalpha()):
            repaired_chars.append(char)
            continue

        left_start = index - 1
        while left_start > 0 and value[left_start - 1].isalpha():
            left_start -= 1
        right_end = index + 1
        while right_end < length - 1 and value[right_end + 1].isalpha():
            right_end += 1
        left_fragment = value[left_start:index]
        right_fragment = value[index + 1 : right_end + 1]

        if (
            len(left_fragment) >= 2
            and len(right_fragment) >= 2
            and left_fragment[-1].islower()
            and right_fragment[0].islower()
        ):
            continue

        repaired_chars.append(char)

    return "".join(repaired_chars)


def _normalize_margin_candidate(line: str) -> str:
    value = _normalize_ocr_text(line)
    value = re.sub(r"\s+", " ", value).strip().lower()
    return value.strip(" -–—_|•·:;,.")


def _is_probable_page_number_line(line: str) -> bool:
    normalized = _normalize_margin_candidate(line)
    if not normalized:
        return False
    compact = re.sub(r"[\s\-–—_.,:;|/\\()[\]{}]+", "", normalized)
    if not compact:
        return False
    if re.fullmatch(r"\d{1,5}", compact):
        return True
    if re.fullmatch(r"(page|pagina|pag|p)\d{1,5}", compact):
        return True
    return re.fullmatch(r"[ivxlcdm]{1,12}", compact) is not None


def _repeated_margin_lines(pages: list[OcrTextPage]) -> set[str]:
    counts: dict[str, int] = {}
    for page in pages:
        lines = [line.strip() for line in _normalize_ocr_text(page.text).splitlines() if line.strip()]
        if not lines:
            continue
        candidates = lines[:3] + lines[-3:]
        seen_in_page: set[str] = set()
        for line in candidates:
            normalized = _normalize_margin_candidate(line)
            if (
                not normalized
                or len(normalized) > 140
                or _is_probable_page_number_line(line)
                or not any(character.isalpha() for character in normalized)
            ):
                continue
            seen_in_page.add(normalized)
        for normalized in seen_in_page:
            counts[normalized] = counts.get(normalized, 0) + 1
    return {line for line, count in counts.items() if count >= 2}


def _strip_margin_artifacts(text: str, repeated_lines: set[str]) -> str:
    original_lines = [line.strip() for line in _normalize_ocr_text(text).splitlines()]
    if not original_lines:
        return ""

    start = 0
    end = len(original_lines)
    while start < end:
        line = original_lines[start]
        normalized = _normalize_margin_candidate(line)
        if not normalized:
            start += 1
            continue
        if normalized in repeated_lines or _is_probable_page_number_line(line):
            start += 1
            continue
        break

    while end > start:
        line = original_lines[end - 1]
        normalized = _normalize_margin_candidate(line)
        if not normalized:
            end -= 1
            continue
        if normalized in repeated_lines or _is_probable_page_number_line(line):
            end -= 1
            continue
        break

    cleaned_lines: list[str] = []
    margin_limit = max(0, end - start)
    for index, line in enumerate(original_lines[start:end]):
        normalized = _normalize_margin_candidate(line)
        is_margin = index < 3 or index >= margin_limit - 3
        if is_margin and normalized in repeated_lines:
            continue
        if _is_probable_page_number_line(line):
            continue
        cleaned_lines.append(line)

    return "\n".join(cleaned_lines).strip()


def _join_text_fragments(left: str, right: str) -> str:
    left_value = left.rstrip()
    right_value = right.lstrip()
    if not left_value:
        return right_value
    if not right_value:
        return left_value
    if left_value.endswith("-") and right_value[:1].isalpha():
        return left_value[:-1] + right_value
    if re.match(r"^[,.;:!?%)\]\}]", right_value):
        return left_value + right_value
    return f"{left_value} {right_value}"


def _page_boundary_requires_merge(left: str, right: str) -> bool:
    left_value = left.rstrip()
    right_value = right.lstrip()
    if not left_value or not right_value:
        return False
    if left_value.endswith("-") and right_value[:1].isalpha():
        return True

    leading = right_value[0]
    if leading.islower() or leading.isdigit() or leading in ",.;:!?%)]}":
        return True
    if leading in "\"'“‘«(" and len(right_value) > 1 and right_value[1].islower():
        return True

    trimmed_left = left_value.rstrip("\"'»”’)]}")
    if not trimmed_left:
        return False
    trailing = trimmed_left[-1]
    if trailing.islower() or trailing.isdigit():
        return True
    return trailing in ",;:([{-–—/"


def _extract_keep_image_marker(text: str) -> tuple[str, bool]:
    preserve_image = False
    cleaned_lines: list[str] = []
    for line in _normalize_ocr_text(text).splitlines():
        if line.strip() == "[[KEEP_PAGE_IMAGE]]":
            preserve_image = True
            continue
        cleaned_lines.append(line)
    return "\n".join(cleaned_lines).strip(), preserve_image


def _paragraphs_from_text(text: str) -> list[str]:
    return [block.strip() for block in re.split(r"\n\s*\n", text) if block.strip()]


def _normalize_ai_paragraph_lines(lines: list[str]) -> str:
    return _normalize_export_text("\n".join(lines)).strip()


def _ai_inline_fragments(text: str) -> list[tuple[str, str]]:
    pattern = re.compile(r"(\*\*[^*\n]+\*\*|\*[^*\n]+\*)")
    fragments: list[tuple[str, str]] = []
    position = 0
    for match in pattern.finditer(text):
        if match.start() > position:
            fragments.append(("text", text[position:match.start()]))
        token = match.group(0)
        if token.startswith("**") and token.endswith("**"):
            fragments.append(("bold", token[2:-2]))
        else:
            fragments.append(("italic", token[1:-1]))
        position = match.end()
    if position < len(text):
        fragments.append(("text", text[position:]))
    return fragments


def _ai_inline_to_html(text: str) -> str:
    chunks: list[str] = []
    for kind, value in _ai_inline_fragments(text):
        escaped = html.escape(value)
        if kind == "bold":
            chunks.append(f"<strong>{escaped}</strong>")
        elif kind == "italic":
            chunks.append(f"<em>{escaped}</em>")
        else:
            chunks.append(escaped)
    return "".join(chunks).replace("\n", "<br/>")


def _markdown_asset_path(target_path: Path, image_path: Path) -> Path:
    assets_dir = ensure_directory(target_path.parent / f"{target_path.stem}_assets")
    candidate = assets_dir / image_path.name
    if candidate != image_path and not candidate.exists():
        shutil.copy2(image_path, candidate)
    return candidate


def _embedded_image_uri(image_path: Path) -> str:
    data = base64.b64encode(image_path.read_bytes()).decode("ascii")
    return f"data:image/png;base64,{data}"


def _scaled_image_size(image_path: Path, max_width: float) -> tuple[float, float]:
    with Image.open(image_path) as image:
        width, height = image.size
    if width <= 0 or height <= 0:
        return max_width, max_width * 0.75
    scale = min(1.0, max_width / float(width))
    return float(width) * scale, float(height) * scale


_HEADING_CONNECTORS = {
    "a",
    "al",
    "and",
    "as",
    "at",
    "by",
    "con",
    "da",
    "de",
    "del",
    "des",
    "di",
    "do",
    "dos",
    "el",
    "en",
    "for",
    "from",
    "la",
    "las",
    "los",
    "of",
    "on",
    "or",
    "para",
    "por",
    "sin",
    "the",
    "to",
    "un",
    "una",
    "y",
}


_KNOWN_HEADING_LABELS = {
    "agradecimiento",
    "agradecimientos",
    "introduccion",
    "conclusion",
    "conclusiones",
    "resumen",
    "sumario",
    "prefacio",
    "prologo",
    "epilogo",
    "bibliografia",
    "referencias",
    "glosario",
    "notas",
    "anexo",
    "anexos",
    "apendice",
    "apendices",
    "indice",
}
_SECTION_HEADING_LABELS = {"capitulo", "chapter", "parte", "part", "seccion", "section"}
_SECTION_HEADING_PATTERN = re.compile(r"^(capitulo|chapter|parte|part|seccion|section)\s+[ivxlcdm\d]+(?:\b.*)?$")


def _heading_words(text: str) -> list[str]:
    return re.findall(r"[A-Za-zÀ-ÿ0-9]+", text)


def _is_titleish_word(word: str) -> bool:
    if not word:
        return False
    if word.isupper():
        return True
    first = word[0]
    return first.isalpha() and first.isupper()


def _looks_like_known_heading(text: str) -> bool:
    candidate = text.strip().strip(":")
    folded = "".join(
        character
        for character in unicodedata.normalize("NFKD", candidate)
        if not unicodedata.combining(character)
    ).lower()
    if folded in _KNOWN_HEADING_LABELS or _SECTION_HEADING_PATTERN.fullmatch(folded) is not None:
        return True
    if any(SequenceMatcher(None, folded, label).ratio() >= 0.88 for label in _KNOWN_HEADING_LABELS):
        return True
    section_match = re.match(r"^([a-z?]+)\s+([ivxlcdm\d]+(?:\b.*)?)$", folded)
    if section_match:
        prefix = section_match.group(1)
        return any(SequenceMatcher(None, prefix, label).ratio() >= 0.84 for label in _SECTION_HEADING_LABELS)
    return False


def _infer_heading_kind(
    line: str,
    *,
    prev_blank: bool,
    next_blank: bool,
    next_nonempty: str,
    at_document_start: bool,
) -> str | None:
    candidate = line.strip().strip(" -–—")
    if len(candidate) < 3 or len(candidate) > 120:
        return None
    if candidate.startswith(("#", ">", "[[")) or re.match(r"^(?:[-*]|\d+\.)\s+", candidate):
        return None
    if candidate.endswith((".", "?", "!", ";")):
        return None
    if candidate.count(",") > 1:
        return None

    if _looks_like_known_heading(candidate):
        return "heading1"

    words = _heading_words(candidate)
    if not words or len(words) > 14:
        return None

    alphabetic_words = [word for word in words if any(character.isalpha() for character in word)]
    if not alphabetic_words:
        return None
    content_words = [word for word in alphabetic_words if word.lower() not in _HEADING_CONNECTORS]
    if not content_words:
        content_words = alphabetic_words

    titleish_ratio = sum(1 for word in content_words if _is_titleish_word(word)) / max(len(content_words), 1)
    compact_candidate = candidate.replace(" ", "")
    is_all_caps = compact_candidate.isupper() and any(character.isalpha() for character in compact_candidate)
    short_isolated_line = len(candidate) <= 70 and len(words) <= 10 and "," not in candidate
    followed_by_body = bool(next_nonempty) and len(next_nonempty) >= len(candidate) + 12

    if is_all_caps and (prev_blank or next_blank):
        return "heading1"
    if titleish_ratio >= 0.7 and (prev_blank or next_blank):
        return "heading1" if at_document_start and len(words) <= 8 else "heading2"
    if short_isolated_line and prev_blank and (next_blank or followed_by_body):
        return "heading2"
    if candidate.endswith(":") and prev_blank:
        return "heading2"
    if titleish_ratio >= 0.45 and prev_blank and followed_by_body:
        return "heading2"
    return None


def _blocks_from_ai_text(text: str) -> list[StructuredOcrBlock]:
    lines = _normalize_ocr_text(text).splitlines()
    blocks: list[StructuredOcrBlock] = []
    paragraph_lines: list[str] = []

    def flush_paragraph() -> None:
        nonlocal paragraph_lines
        if not paragraph_lines:
            return
        paragraph = _normalize_ai_paragraph_lines(paragraph_lines)
        if paragraph:
            blocks.append(StructuredOcrBlock("paragraph", paragraph))
        paragraph_lines = []

    for index, raw_line in enumerate(lines):
        line = raw_line.strip()
        if not line:
            flush_paragraph()
            continue
        prev_blank = index == 0 or not lines[index - 1].strip()
        next_blank = index >= len(lines) - 1 or not lines[index + 1].strip()
        next_nonempty = ""
        for future_index in range(index + 1, len(lines)):
            candidate = lines[future_index].strip()
            if candidate:
                next_nonempty = candidate
                break
        if line.startswith("### "):
            flush_paragraph()
            blocks.append(StructuredOcrBlock("heading3", line[4:].strip()))
            continue
        if line.startswith("## "):
            flush_paragraph()
            blocks.append(StructuredOcrBlock("heading2", line[3:].strip()))
            continue
        if line.startswith("# "):
            flush_paragraph()
            blocks.append(StructuredOcrBlock("heading1", line[2:].strip()))
            continue
        if line.startswith("> "):
            flush_paragraph()
            blocks.append(StructuredOcrBlock("quote", _normalize_ai_paragraph_lines([line[2:].strip()])))
            continue
        if re.match(r"^(?:[-*]|\d+\.)\s+", line):
            flush_paragraph()
            item_text = re.sub(r"^(?:[-*]|\d+\.)\s+", "", line, count=1).strip()
            blocks.append(StructuredOcrBlock("list_item", _normalize_ai_paragraph_lines([item_text])))
            continue
        inferred_heading_kind = _infer_heading_kind(
            line,
            prev_blank=prev_blank,
            next_blank=next_blank,
            next_nonempty=next_nonempty,
            at_document_start=not blocks and not paragraph_lines,
        )
        if inferred_heading_kind is not None:
            flush_paragraph()
            blocks.append(StructuredOcrBlock(inferred_heading_kind, line))
            continue
        paragraph_lines.append(line)

    flush_paragraph()
    return blocks


def _build_ollama_document_blocks(
    pages: list[OcrTextPage],
    progress_callback: ProgressCallback | None = None,
) -> list[StructuredOcrBlock]:
    if not pages:
        return []
    repeated_lines = _repeated_margin_lines(pages)
    if progress_callback is not None:
        progress_callback(
            len(pages),
            len(pages),
            "Cleaning repeated headers, footers, page numbers, and visual markers from AI OCR output...",
        )

    merged_blocks: list[StructuredOcrBlock] = []
    for page in pages:
        cleaned_page_text = _strip_margin_artifacts(page.text, repeated_lines)
        cleaned_page_text, preserve_image = _extract_keep_image_marker(cleaned_page_text)
        normalized_page_text = _normalize_export_text(cleaned_page_text)
        page_blocks = [StructuredOcrBlock("paragraph", paragraph) for paragraph in _paragraphs_from_text(normalized_page_text)]

        if not page_blocks:
            if preserve_image:
                merged_blocks.append(StructuredOcrBlock("image", image_path=page.image_path))
            continue

        if (
            merged_blocks
            and merged_blocks[-1].kind == "paragraph"
            and page_blocks[0].kind == "paragraph"
            and _page_boundary_requires_merge(merged_blocks[-1].text, page_blocks[0].text)
        ):
            merged_blocks[-1] = StructuredOcrBlock(
                "paragraph",
                _join_text_fragments(merged_blocks[-1].text, page_blocks[0].text),
            )
            remaining_blocks = page_blocks[1:]
        else:
            remaining_blocks = page_blocks
        if preserve_image:
            merged_blocks.append(StructuredOcrBlock("image", image_path=page.image_path))
        merged_blocks.extend(remaining_blocks)
    return merged_blocks


def _build_ollama_continuous_text(
    pages: list[OcrTextPage],
    progress_callback: ProgressCallback | None = None,
) -> str:
    paragraphs = [block.text for block in _build_ollama_document_blocks(pages, progress_callback) if block.kind == "paragraph"]
    return "\n\n".join(paragraphs).strip()


def _write_ollama_text_txt(target_path: Path, pages: list[OcrTextPage]) -> None:
    chunks: list[str] = []
    for page in pages:
        chunks.append(f"===== Page {page.page_number} =====\n{_normalize_export_text(page.text)}")
    target_path.write_text("\n\n".join(chunks) + "\n", encoding="utf-8")


def _write_ollama_text_markdown(target_path: Path, pages: list[OcrTextPage]) -> None:
    title = _ocr_text_document_title(target_path)
    chunks = [f"# {title}"]
    for page in pages:
        chunks.append(f"## Page {page.page_number}\n\n```text\n{_normalize_export_text(page.text)}\n```")
    target_path.write_text("\n\n".join(chunks) + "\n", encoding="utf-8")


def _write_ollama_text_html(target_path: Path, pages: list[OcrTextPage]) -> None:
    title = html.escape(_ocr_text_document_title(target_path))
    sections = []
    for page in pages:
        sections.append(
            "<section class=\"page\">"
            f"<h2>Page {page.page_number}</h2>"
            f"<pre>{html.escape(_normalize_export_text(page.text))}</pre>"
            "</section>"
        )
    document = (
        "<!DOCTYPE html><html><head><meta charset=\"utf-8\">"
        f"<title>{title}</title>"
        "<style>"
        "body{font-family:Segoe UI,Arial,sans-serif;margin:24px;line-height:1.4;}"
        "section.page{margin-bottom:32px;padding-bottom:24px;border-bottom:1px solid #d0d0d0;}"
        "pre{white-space:pre-wrap;word-wrap:break-word;font-family:Consolas,'Courier New',monospace;"
        "background:#f6f6f6;padding:16px;border-radius:8px;}"
        "</style></head><body>"
        f"<h1>{title}</h1>"
        f"{''.join(sections)}"
        "</body></html>"
    )
    target_path.write_text(document, encoding="utf-8")


def _write_ollama_text_pdf(target_path: Path, pages: list[OcrTextPage]) -> None:
    ensure_directory(target_path.parent)
    document = SimpleDocTemplate(
        str(target_path),
        pagesize=A4,
        leftMargin=48,
        rightMargin=48,
        topMargin=56,
        bottomMargin=56,
        title=_ocr_text_document_title(target_path),
    )
    styles = getSampleStyleSheet()
    body_style = ParagraphStyle(
        "OcrBody",
        parent=styles["BodyText"],
        fontName="Times-Roman",
        fontSize=11.5,
        leading=15,
        alignment=TA_JUSTIFY,
        spaceAfter=10,
    )
    story: list[object] = []
    for page_index, page in enumerate(pages):
        blocks = [block.strip() for block in re.split(r"\n\s*\n", _normalize_export_text(page.text)) if block.strip()]
        for block in blocks:
            paragraph_html = html.escape(block).replace("\n", "<br/>")
            story.append(Paragraph(paragraph_html, body_style))
            story.append(Spacer(1, 8))
        if page_index < len(pages) - 1:
            story.append(PageBreak())
    if not story:
        story.append(Paragraph("", body_style))
    document.build(story)


def _write_ollama_text_docx(target_path: Path, pages: list[OcrTextPage]) -> None:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError(
            "DOCX export requires python-docx. Install dependencies with: .\\venv\\Scripts\\pip.exe install python-docx"
        ) from exc

    document = Document()
    document.core_properties.title = _ocr_text_document_title(target_path)
    document.add_heading(_ocr_text_document_title(target_path), level=0)
    for page in pages:
        document.add_heading(f"Page {page.page_number}", level=1)
        blocks = re.split(r"\n\s*\n", _normalize_export_text(page.text))
        for block in blocks:
            paragraph = document.add_paragraph()
            lines = block.splitlines() or [""]
            for index, line in enumerate(lines):
                paragraph.add_run(line)
                if index < len(lines) - 1:
                    paragraph.add_run().add_break()
    document.save(str(target_path))


def _write_ollama_text_epub(target_path: Path, pages: list[OcrTextPage]) -> None:
    try:
        from ebooklib import epub
    except ImportError as exc:
        raise RuntimeError(
            "EPUB export requires ebooklib. Install dependencies with: .\\venv\\Scripts\\pip.exe install ebooklib"
        ) from exc

    title = _ocr_text_document_title(target_path)
    book = epub.EpubBook()
    book.set_identifier(f"ust-{int(time.time())}")
    book.set_title(title)
    book.set_language("en")

    style = (
        "body{font-family:serif;line-height:1.4;}"
        "pre{white-space:pre-wrap;font-family:monospace;}"
    )
    css_item = epub.EpubItem(
        uid="style-default",
        file_name="styles/default.css",
        media_type="text/css",
        content=style.encode("utf-8"),
    )
    book.add_item(css_item)

    chapters = []
    for page in pages:
        chapter = epub.EpubHtml(
            title=f"Page {page.page_number}",
            file_name=f"page_{page.page_number:04d}.xhtml",
            lang="en",
        )
        chapter.content = (
            "<html><head><link rel=\"stylesheet\" href=\"styles/default.css\" type=\"text/css\"/></head><body>"
            f"<h1>Page {page.page_number}</h1>"
            f"<pre>{html.escape(_normalize_export_text(page.text))}</pre>"
            "</body></html>"
        )
        book.add_item(chapter)
        chapters.append(chapter)

    book.toc = tuple(chapters)
    book.spine = ["nav", *chapters]
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())
    epub.write_epub(str(target_path), book, {})


def _write_continuous_text_txt(target_path: Path, text: str) -> None:
    target_path.write_text(text.rstrip() + "\n", encoding="utf-8")


def _write_continuous_text_markdown(target_path: Path, text: str) -> None:
    title = _ocr_text_document_title(target_path)
    body = text.strip()
    content = f"# {title}\n\n{body}\n" if body else f"# {title}\n"
    target_path.write_text(content, encoding="utf-8")


def _write_continuous_text_html(target_path: Path, text: str) -> None:
    title = html.escape(_ocr_text_document_title(target_path))
    paragraphs = _paragraphs_from_text(text)
    body_html = "".join(
        f"<p>{html.escape(paragraph).replace(chr(10), '<br/>')}</p>" for paragraph in paragraphs
    ) or "<p></p>"
    document = (
        "<!DOCTYPE html><html><head><meta charset=\"utf-8\">"
        f"<title>{title}</title>"
        "<style>"
        "body{font-family:Segoe UI,Arial,sans-serif;margin:24px;line-height:1.5;max-width:900px;}"
        "p{margin:0 0 1em 0;}"
        "</style></head><body>"
        f"<h1>{title}</h1>"
        f"{body_html}"
        "</body></html>"
    )
    target_path.write_text(document, encoding="utf-8")


def _write_continuous_text_pdf(target_path: Path, text: str) -> None:
    ensure_directory(target_path.parent)
    document = SimpleDocTemplate(
        str(target_path),
        pagesize=A4,
        leftMargin=48,
        rightMargin=48,
        topMargin=56,
        bottomMargin=56,
        title=_ocr_text_document_title(target_path),
    )
    styles = getSampleStyleSheet()
    body_style = ParagraphStyle(
        "OcrContinuousBody",
        parent=styles["BodyText"],
        fontName="Times-Roman",
        fontSize=11.5,
        leading=15,
        alignment=TA_JUSTIFY,
        spaceAfter=10,
    )
    story: list[object] = []
    for paragraph in _paragraphs_from_text(text):
        paragraph_html = html.escape(paragraph).replace("\n", "<br/>")
        story.append(Paragraph(paragraph_html, body_style))
        story.append(Spacer(1, 8))
    if not story:
        story.append(Paragraph("", body_style))
    document.build(story)


def _write_continuous_text_docx(target_path: Path, text: str) -> None:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError(
            "DOCX export requires python-docx. Install dependencies with: .\\venv\\Scripts\\pip.exe install python-docx"
        ) from exc

    document = Document()
    document.core_properties.title = _ocr_text_document_title(target_path)
    document.add_heading(_ocr_text_document_title(target_path), level=0)
    for paragraph_text in _paragraphs_from_text(text):
        paragraph = document.add_paragraph()
        lines = paragraph_text.splitlines() or [""]
        for index, line in enumerate(lines):
            paragraph.add_run(line)
            if index < len(lines) - 1:
                paragraph.add_run().add_break()
    document.save(str(target_path))


def _write_continuous_text_epub(target_path: Path, text: str) -> None:
    try:
        from ebooklib import epub
    except ImportError as exc:
        raise RuntimeError(
            "EPUB export requires ebooklib. Install dependencies with: .\\venv\\Scripts\\pip.exe install ebooklib"
        ) from exc

    title = _ocr_text_document_title(target_path)
    book = epub.EpubBook()
    book.set_identifier(f"ust-{int(time.time())}")
    book.set_title(title)
    book.set_language("en")

    style = "body{font-family:serif;line-height:1.5;} p{margin:0 0 1em 0;}"
    css_item = epub.EpubItem(
        uid="style-default",
        file_name="styles/default.css",
        media_type="text/css",
        content=style.encode("utf-8"),
    )
    book.add_item(css_item)

    paragraphs = "".join(
        f"<p>{html.escape(paragraph).replace(chr(10), '<br/>')}</p>" for paragraph in _paragraphs_from_text(text)
    ) or "<p></p>"
    chapter = epub.EpubHtml(title=title, file_name="content.xhtml", lang="en")
    chapter.content = (
        "<html><head><link rel=\"stylesheet\" href=\"styles/default.css\" type=\"text/css\"/></head><body>"
        f"<h1>{html.escape(title)}</h1>"
        f"{paragraphs}"
        "</body></html>"
    )
    book.add_item(chapter)
    book.toc = (chapter,)
    book.spine = ["nav", chapter]
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())
    epub.write_epub(str(target_path), book, {})


def _blocks_to_plain_text(blocks: list[StructuredOcrBlock]) -> str:
    parts: list[str] = []
    for block in blocks:
        if block.kind == "image":
            continue
        if block.kind == "list_item":
            parts.append(f"- {block.text}")
        elif block.kind == "quote":
            parts.append(f"> {block.text}")
        else:
            parts.append(block.text)
    return "\n\n".join(part for part in parts if part.strip()).strip()


def _write_ai_structured_txt(target_path: Path, blocks: list[StructuredOcrBlock]) -> None:
    _write_continuous_text_txt(target_path, _blocks_to_plain_text(blocks))


def _write_ai_structured_markdown(target_path: Path, blocks: list[StructuredOcrBlock]) -> None:
    lines = [f"# {_ocr_text_document_title(target_path)}", ""]
    for block in blocks:
        if block.kind == "image" and block.image_path is not None:
            asset_path = _markdown_asset_path(target_path, block.image_path)
            relative_path = asset_path.relative_to(target_path.parent).as_posix()
            lines.extend([f"![Preserved page image]({relative_path})", ""])
        elif block.kind == "heading1":
            lines.extend([f"# {block.text}", ""])
        elif block.kind == "heading2":
            lines.extend([f"## {block.text}", ""])
        elif block.kind == "heading3":
            lines.extend([f"### {block.text}", ""])
        elif block.kind == "quote":
            lines.extend([f"> {block.text}", ""])
        elif block.kind == "list_item":
            lines.extend([f"- {block.text}", ""])
        else:
            lines.extend([block.text, ""])
    target_path.write_text("\n".join(lines).rstrip() + "\n", encoding="utf-8")


def _write_ai_structured_html(target_path: Path, blocks: list[StructuredOcrBlock]) -> None:
    title = html.escape(_ocr_text_document_title(target_path))
    body_parts: list[str] = []
    index = 0
    while index < len(blocks):
        block = blocks[index]
        if block.kind == "image" and block.image_path is not None:
            body_parts.append(
                f"<figure><img src=\"{_embedded_image_uri(block.image_path)}\" alt=\"Preserved page image\" /></figure>"
            )
        elif block.kind == "heading1":
            body_parts.append(f"<h1>{_ai_inline_to_html(block.text)}</h1>")
        elif block.kind == "heading2":
            body_parts.append(f"<h2>{_ai_inline_to_html(block.text)}</h2>")
        elif block.kind == "heading3":
            body_parts.append(f"<h3>{_ai_inline_to_html(block.text)}</h3>")
        elif block.kind == "quote":
            body_parts.append(f"<blockquote>{_ai_inline_to_html(block.text)}</blockquote>")
        elif block.kind == "list_item":
            items: list[str] = []
            while index < len(blocks) and blocks[index].kind == "list_item":
                items.append(f"<li>{_ai_inline_to_html(blocks[index].text)}</li>")
                index += 1
            body_parts.append(f"<ul>{''.join(items)}</ul>")
            continue
        else:
            body_parts.append(f"<p>{_ai_inline_to_html(block.text)}</p>")
        index += 1

    document = (
        "<!DOCTYPE html><html><head><meta charset=\"utf-8\">"
        f"<title>{title}</title>"
        "<style>"
        "body{font-family:Segoe UI,Arial,sans-serif;margin:24px;line-height:1.55;max-width:900px;}"
        "figure{margin:1.2em 0;}img{max-width:100%;height:auto;border-radius:8px;}"
        "blockquote{margin:1em 0;padding-left:1em;border-left:3px solid #ccc;color:#444;font-style:italic;}"
        "ul{margin:0.75em 0 0.75em 1.5em;}p{margin:0 0 1em 0;}"
        "</style></head><body>"
        f"<h1>{title}</h1>"
        f"{''.join(body_parts)}"
        "</body></html>"
    )
    target_path.write_text(document, encoding="utf-8")


def _write_ai_structured_pdf(target_path: Path, blocks: list[StructuredOcrBlock]) -> None:
    ensure_directory(target_path.parent)
    document = SimpleDocTemplate(
        str(target_path),
        pagesize=A4,
        leftMargin=48,
        rightMargin=48,
        topMargin=56,
        bottomMargin=56,
        title=_ocr_text_document_title(target_path),
    )
    styles = getSampleStyleSheet()
    heading1_style = ParagraphStyle("AiHeading1", parent=styles["Heading1"], spaceAfter=12)
    heading2_style = ParagraphStyle("AiHeading2", parent=styles["Heading2"], spaceAfter=10)
    heading3_style = ParagraphStyle("AiHeading3", parent=styles["Heading3"], spaceAfter=8)
    body_style = ParagraphStyle(
        "AiBody",
        parent=styles["BodyText"],
        fontName="Times-Roman",
        fontSize=11.5,
        leading=15,
        alignment=TA_JUSTIFY,
        spaceAfter=10,
    )
    quote_style = ParagraphStyle("AiQuote", parent=body_style, leftIndent=20, rightIndent=12, italic=True, textColor="#444444")
    list_style = ParagraphStyle("AiList", parent=body_style, leftIndent=18, bulletIndent=6)
    story: list[object] = []
    max_width = A4[0] - 96
    for block in blocks:
        if block.kind == "image" and block.image_path is not None:
            width, height = _scaled_image_size(block.image_path, max_width)
            story.append(PlatypusImage(str(block.image_path), width=width, height=height))
            story.append(Spacer(1, 10))
        elif block.kind == "heading1":
            story.append(Paragraph(_ai_inline_to_html(block.text), heading1_style))
        elif block.kind == "heading2":
            story.append(Paragraph(_ai_inline_to_html(block.text), heading2_style))
        elif block.kind == "heading3":
            story.append(Paragraph(_ai_inline_to_html(block.text), heading3_style))
        elif block.kind == "quote":
            story.append(Paragraph(_ai_inline_to_html(block.text), quote_style))
        elif block.kind == "list_item":
            story.append(Paragraph(f"• {_ai_inline_to_html(block.text)}", list_style))
        else:
            story.append(Paragraph(_ai_inline_to_html(block.text), body_style))
        story.append(Spacer(1, 6))
    if not story:
        story.append(Paragraph("", body_style))
    document.build(story)


def _docx_add_inline_runs(paragraph, text: str) -> None:
    for kind, value in _ai_inline_fragments(text):
        run = paragraph.add_run(value)
        if kind == "bold":
            run.bold = True
        elif kind == "italic":
            run.italic = True


def _write_ai_structured_docx(target_path: Path, blocks: list[StructuredOcrBlock]) -> None:
    try:
        from docx import Document
        from docx.shared import Inches
    except ImportError as exc:
        raise RuntimeError(
            "DOCX export requires python-docx. Install dependencies with: .\\venv\\Scripts\\pip.exe install python-docx"
        ) from exc

    document = Document()
    document.core_properties.title = _ocr_text_document_title(target_path)
    document.add_heading(_ocr_text_document_title(target_path), level=0)
    for block in blocks:
        if block.kind == "image" and block.image_path is not None:
            document.add_picture(str(block.image_path), width=Inches(6.0))
            continue
        if block.kind == "heading1":
            paragraph = document.add_heading("", level=1)
        elif block.kind == "heading2":
            paragraph = document.add_heading("", level=2)
        elif block.kind == "heading3":
            paragraph = document.add_heading("", level=3)
        elif block.kind == "list_item":
            paragraph = document.add_paragraph(style="List Bullet")
        else:
            paragraph = document.add_paragraph()
            if block.kind == "quote":
                paragraph.paragraph_format.left_indent = Inches(0.35)
        _docx_add_inline_runs(paragraph, block.text)
    document.save(str(target_path))


def _write_ai_structured_epub(target_path: Path, blocks: list[StructuredOcrBlock]) -> None:
    try:
        from ebooklib import epub
    except ImportError as exc:
        raise RuntimeError(
            "EPUB export requires ebooklib. Install dependencies with: .\\venv\\Scripts\\pip.exe install ebooklib"
        ) from exc

    title = _ocr_text_document_title(target_path)
    book = epub.EpubBook()
    book.set_identifier(f"ust-{int(time.time())}")
    book.set_title(title)
    book.set_language("en")
    style = (
        "body{font-family:serif;line-height:1.5;} p{margin:0 0 1em 0;}"
        "blockquote{margin:1em 0;padding-left:1em;border-left:3px solid #ccc;font-style:italic;}"
        "img{max-width:100%;height:auto;display:block;margin:1em 0;}"
    )
    css_item = epub.EpubItem(
        uid="style-default",
        file_name="styles/default.css",
        media_type="text/css",
        content=style.encode("utf-8"),
    )
    book.add_item(css_item)

    body_parts: list[str] = []
    image_count = 0
    index = 0
    while index < len(blocks):
        block = blocks[index]
        if block.kind == "image" and block.image_path is not None:
            image_count += 1
            image_name = f"images/image_{image_count:04d}.png"
            image_item = epub.EpubItem(
                uid=f"image-{image_count}",
                file_name=image_name,
                media_type="image/png",
                content=block.image_path.read_bytes(),
            )
            book.add_item(image_item)
            body_parts.append(f"<figure><img src=\"{image_name}\" alt=\"Preserved page image\"/></figure>")
        elif block.kind == "heading1":
            body_parts.append(f"<h1>{_ai_inline_to_html(block.text)}</h1>")
        elif block.kind == "heading2":
            body_parts.append(f"<h2>{_ai_inline_to_html(block.text)}</h2>")
        elif block.kind == "heading3":
            body_parts.append(f"<h3>{_ai_inline_to_html(block.text)}</h3>")
        elif block.kind == "quote":
            body_parts.append(f"<blockquote>{_ai_inline_to_html(block.text)}</blockquote>")
        elif block.kind == "list_item":
            items: list[str] = []
            while index < len(blocks) and blocks[index].kind == "list_item":
                items.append(f"<li>{_ai_inline_to_html(blocks[index].text)}</li>")
                index += 1
            body_parts.append(f"<ul>{''.join(items)}</ul>")
            continue
        else:
            body_parts.append(f"<p>{_ai_inline_to_html(block.text)}</p>")
        index += 1

    chapter = epub.EpubHtml(title=title, file_name="content.xhtml", lang="en")
    chapter.content = (
        "<html><head><link rel=\"stylesheet\" href=\"styles/default.css\" type=\"text/css\"/></head><body>"
        f"<h1>{html.escape(title)}</h1>"
        f"{''.join(body_parts)}"
        "</body></html>"
    )
    book.add_item(chapter)
    book.toc = (chapter,)
    book.spine = ["nav", chapter]
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())
    epub.write_epub(str(target_path), book, {})


def export_ollama_ocr_text_document(
    image_paths: Iterable[Path],
    target_path: Path,
    base_url: str,
    model_name: str,
    output_format: str,
    progress_callback: ProgressCallback | None = None,
    prompt: str = DEFAULT_OCR_PROMPT,
) -> Path:
    normalized_format = output_format.strip().lower().lstrip(".")
    if normalized_format not in {"txt", "md", "markdown", "html", "epub", "docx"}:
        raise RuntimeError(f"Unsupported OCR export format: {output_format}")

    ensure_directory(target_path.parent)
    pages = _collect_ollama_ocr_pages(
        image_paths=image_paths,
        base_url=base_url,
        model_name=model_name,
        progress_callback=progress_callback,
        prompt=prompt,
    )
    document_blocks = _build_ollama_document_blocks(pages, progress_callback=progress_callback)

    if progress_callback is not None:
        progress_callback(len(pages), len(pages), f"Writing {normalized_format.upper()} OCR document...")

    if normalized_format == "txt":
        _write_ai_structured_txt(target_path, document_blocks)
    elif normalized_format in {"md", "markdown"}:
        _write_ai_structured_markdown(target_path, document_blocks)
    elif normalized_format == "html":
        _write_ai_structured_html(target_path, document_blocks)
    elif normalized_format == "epub":
        _write_ai_structured_epub(target_path, document_blocks)
    else:
        _write_ai_structured_docx(target_path, document_blocks)

    if progress_callback is not None:
        progress_callback(len(pages), len(pages), f"Ollama OCR {normalized_format.upper()} export complete.")
    return target_path


def export_surya_ocr_text_document(
    image_paths: Iterable[Path],
    target_path: Path,
    output_format: str,
    progress_callback: ProgressCallback | None = None,
) -> Path:
    normalized_format = output_format.strip().lower().lstrip(".")
    if normalized_format not in {"txt", "md", "markdown", "html", "epub", "docx"}:
        raise RuntimeError(f"Unsupported OCR export format: {output_format}")

    ensure_directory(target_path.parent)
    pages = _collect_surya_ocr_pages(
        image_paths=image_paths,
        progress_callback=progress_callback,
    )

    if progress_callback is not None:
        progress_callback(len(pages), len(pages), f"Writing {normalized_format.upper()} Surya OCR document...")

    if normalized_format == "txt":
        _write_ollama_text_txt(target_path, pages)
    elif normalized_format in {"md", "markdown"}:
        _write_ollama_text_markdown(target_path, pages)
    elif normalized_format == "html":
        _write_ollama_text_html(target_path, pages)
    elif normalized_format == "epub":
        _write_ollama_text_epub(target_path, pages)
    else:
        _write_ollama_text_docx(target_path, pages)

    if progress_callback is not None:
        progress_callback(len(pages), len(pages), f"Surya OCR {normalized_format.upper()} export complete.")
    return target_path


def export_ollama_ocr_text_pdf(
    image_paths: Iterable[Path],
    pdf_path: Path,
    base_url: str,
    model_name: str,
    progress_callback: ProgressCallback | None = None,
    prompt: str = DEFAULT_OCR_PROMPT,
) -> Path:
    ensure_directory(pdf_path.parent)
    pages = _collect_ollama_ocr_pages(
        image_paths=image_paths,
        base_url=base_url,
        model_name=model_name,
        progress_callback=progress_callback,
        prompt=prompt,
    )
    document_blocks = _build_ollama_document_blocks(pages, progress_callback=progress_callback)
    if progress_callback is not None:
        progress_callback(len(pages), len(pages), "Writing AI OCR PDF text document...")
    _write_ai_structured_pdf(pdf_path, document_blocks)
    if not pdf_has_extractable_text(pdf_path):
        raise RuntimeError("AI OCR PDF export finished, but the exported PDF does not contain selectable text.")
    if progress_callback is not None:
        progress_callback(len(pages), len(pages), "AI OCR PDF export complete. Searchable text verified.")
    return pdf_path


def export_surya_ocr_text_pdf(
    image_paths: Iterable[Path],
    pdf_path: Path,
    progress_callback: ProgressCallback | None = None,
) -> Path:
    ensure_directory(pdf_path.parent)
    pages = _collect_surya_ocr_pages(
        image_paths=image_paths,
        progress_callback=progress_callback,
    )
    if progress_callback is not None:
        progress_callback(len(pages), len(pages), "Writing Surya OCR PDF text document...")
    _write_ollama_text_pdf(pdf_path, pages)
    if not pdf_has_extractable_text(pdf_path):
        raise RuntimeError("Surya OCR PDF export finished, but the exported PDF does not contain selectable text.")
    if progress_callback is not None:
        progress_callback(len(pages), len(pages), "Surya OCR PDF export complete. Searchable text verified.")
    return pdf_path


def export_searchable_pdf_ocr(
    image_paths: Iterable[Path],
    pdf_path: Path,
    languages: str = "eng",
    progress_callback: ProgressCallback | None = None,
) -> Path:
    ordered_paths = [Path(path) for path in image_paths]
    if not ordered_paths:
        raise ValueError("No screenshots are available to export.")

    tesseract = find_tesseract_executable()
    if not tesseract:
        raise RuntimeError(
            "Tesseract OCR is not installed or not in PATH. "
            f"Install it and try again. Suggested command: {tesseract_install_command()}"
        )

    available_languages = list_tesseract_languages()
    language_value, _missing_languages = resolve_tesseract_languages(languages, available_languages)
    if not language_value:
        raise RuntimeError("No valid Tesseract languages are available for searchable PDF export.")

    if not ocrmypdf_available():
        if progress_callback is not None:
            progress_callback(
                0,
                len(ordered_paths),
                "OCRmyPDF is not available on this system. Falling back to direct Tesseract PDF OCR.",
            )
        return export_tesseract_ocr_pdf(ordered_paths, pdf_path, languages=language_value, progress_callback=progress_callback)

    ensure_directory(pdf_path.parent)
    with tempfile.TemporaryDirectory(prefix="ust_ocrmypdf_") as temp_dir_name:
        temp_dir = Path(temp_dir_name)
        input_pdf = temp_dir / "input_images.pdf"
        export_pdf(ordered_paths, input_pdf)

        if progress_callback is not None:
            progress_callback(0, len(ordered_paths), "Running OCRmyPDF searchable PDF export...")

        ocrmypdf_env = tesseract_environment()
        path_entries = [ocrmypdf_env.get("PATH", "")]
        tesseract_dir = Path(tesseract).resolve().parent
        ghostscript = find_ghostscript_executable()
        if ghostscript:
            path_entries.insert(0, str(Path(ghostscript).resolve().parent))
        path_entries.insert(0, str(tesseract_dir))
        ocrmypdf_env["PATH"] = os.pathsep.join(entry for entry in path_entries if entry)

        result = subprocess.run(
            [
                sys.executable,
                "-m",
                "ocrmypdf",
                "--force-ocr",
                "--output-type",
                "pdf",
                "--optimize",
                "0",
                "--language",
                language_value,
                str(input_pdf),
                str(pdf_path),
            ],
            capture_output=True,
            text=False,
            check=False,
            env=ocrmypdf_env,
        )
        if result.returncode != 0:
            if progress_callback is not None:
                progress_callback(
                    0,
                    len(ordered_paths),
                    "OCRmyPDF failed. Falling back to direct Tesseract PDF OCR.",
                )
            return export_tesseract_ocr_pdf(ordered_paths, pdf_path, languages=language_value, progress_callback=progress_callback)

    if not pdf_has_extractable_text(pdf_path):
        raise RuntimeError("Searchable PDF export finished, but the exported PDF does not contain selectable text.")
    if progress_callback is not None:
        progress_callback(len(ordered_paths), len(ordered_paths), "Searchable PDF OCR complete. Searchable text verified.")
    return pdf_path
